"""Сервис анонимизации текстов и документов."""

from __future__ import annotations

import io
from uuid import uuid4

from app.config import Settings, get_settings
from app.core.audit_logger import AuditLogger, TimedOperation
from app.core.detector import DetectorOrchestrator
from app.core.encryption import Encryptor
from app.core.masker import MaskingService
from app.db.repositories.audit_repo import AuditRepository
from app.db.repositories.mapping_repo import MappingRepository
from app.models.entity import EntityType, MaskingStrategy
from app.models.request import AnonymizeBatchRequest, AnonymizeTextRequest
from app.models.response import AnonymizeTextResponse, BatchResultItem


class AnonymizeService:
    def __init__(
        self,
        settings: Settings | None = None,
        orchestrator: DetectorOrchestrator | None = None,
        masking: MaskingService | None = None,
        audit_logger: AuditLogger | None = None,
    ) -> None:
        self.settings = settings or get_settings()
        encryptor = Encryptor(self.settings.anonymizer_pseudonym_secret_key)
        mapping_repo = MappingRepository(encryptor=encryptor)
        self.orchestrator = orchestrator or DetectorOrchestrator(self.settings)
        self.masking = masking or MaskingService(self.settings, mapping_repo=mapping_repo)
        self.audit_logger = audit_logger or AuditLogger(self.settings)

    def process_text(self, request: AnonymizeTextRequest) -> AnonymizeTextResponse:
        timer = TimedOperation()
        strategy = request.strategy or MaskingStrategy[self.settings.anonymizer_strategy]
        context_id = str(uuid4())

        entities = self.orchestrator.detect(request.text, request.entity_types)
        anonymized, found = self.masking.apply(
            request.text,
            entities,
            strategy=strategy,
            dry_run=request.dry_run,
            context_id=context_id,
        )

        audit = self.audit_logger.build_log(
            request.text,
            entities,
            strategy,
            timer.elapsed_ms,
            operator_id=request.operator_id,
            legal_basis=request.legal_basis,
            dry_run=request.dry_run,
        )

        return AnonymizeTextResponse(
            anonymized_text=anonymized if not request.dry_run else request.text,
            entities_found=found if request.return_entities else [],
            processing_time_ms=timer.elapsed_ms,
            audit_id=audit.audit_id,
            dry_run=request.dry_run,
        )

    async def process_text_with_audit(
        self,
        request: AnonymizeTextRequest,
        audit_repo: AuditRepository,
    ) -> AnonymizeTextResponse:
        response = self.process_text(request)
        strategy = request.strategy or MaskingStrategy[self.settings.anonymizer_strategy]
        audit = self.audit_logger.build_log(
            request.text,
            self.orchestrator.detect(request.text, request.entity_types),
            strategy,
            response.processing_time_ms,
            operator_id=request.operator_id,
            legal_basis=request.legal_basis,
            dry_run=request.dry_run,
        )
        audit.audit_id = response.audit_id
        await self.audit_logger.persist(audit_repo, audit)
        return response

    def process_batch_item(
        self,
        index: int,
        text: str,
        strategy: MaskingStrategy | None,
        entity_types: list[EntityType] | None,
        dry_run: bool,
        return_entities: bool,
    ) -> BatchResultItem:
        try:
            req = AnonymizeTextRequest(
                text=text,
                strategy=strategy,
                entity_types=entity_types,
                return_entities=return_entities,
                dry_run=dry_run,
            )
            result = self.process_text(req)
            return BatchResultItem(
                index=index,
                anonymized_text=result.anonymized_text,
                entities_found=result.entities_found,
            )
        except Exception as exc:
            return BatchResultItem(index=index, anonymized_text=text, error=str(exc))

    def process_batch(self, request: AnonymizeBatchRequest) -> list[BatchResultItem]:
        strategy = request.strategy
        return [
            self.process_batch_item(
                i,
                text,
                strategy,
                request.entity_types,
                request.dry_run,
                request.return_entities,
            )
            for i, text in enumerate(request.texts)
        ]

    def extract_text_from_file(self, filename: str, content: bytes) -> str:
        lower = filename.lower()
        if lower.endswith(".txt"):
            return content.decode("utf-8", errors="replace")
        if lower.endswith(".docx"):
            from docx import Document

            doc = Document(io.BytesIO(content))
            return "\n".join(p.text for p in doc.paragraphs)
        if lower.endswith(".pdf"):
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(content))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        raise ValueError("Поддерживаются форматы: TXT, DOCX, PDF")

    def build_anonymized_file(self, filename: str, text: str) -> tuple[str, bytes, str]:
        lower = filename.lower()
        if lower.endswith(".txt"):
            return filename, text.encode("utf-8"), "text/plain; charset=utf-8"
        if lower.endswith(".docx"):
            from docx import Document

            doc = Document()
            for line in text.splitlines():
                doc.add_paragraph(line)
            buf = io.BytesIO()
            doc.save(buf)
            return f"anonymized_{filename}", buf.getvalue(), (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        if lower.endswith(".pdf"):
            return f"anonymized_{PathStem(filename)}.txt", text.encode("utf-8"), "text/plain; charset=utf-8"
        raise ValueError("Неподдерживаемый формат файла")


def PathStem(filename: str) -> str:
    return filename.rsplit(".", 1)[0]
